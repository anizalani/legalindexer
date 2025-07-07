
import json
from collections import defaultdict
from typing import Dict, List
import csv
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches
from docx.oxml.section import CT_SectPr
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from lxml import etree
import re

def get_table_of_contents(pdf_path: str, page_offset: int) -> Dict:
    """Extract a table of contents from the PDF based on font size."""
    doc = fitz.open(pdf_path)
    toc = {}
    current_headings = [None, None, None, None]

    for page_num, page in enumerate(doc):
        blocks = page.get_text("dict")
        for block in blocks["blocks"]:
            if "lines" in block:
                for line in block["lines"]:
                    if "spans" in line:
                        for span in line["spans"]:
                            font_size = round(span["size"])
                            text = span["text"].strip()
                            if text and not text.isnumeric():
                                # Clean statutory references from headings
                                text = re.sub(r'§\s*\d+(\.\d+)?', '', text).strip()
                                
                                # Level 1: ALL CAPS, bold/heavy font
                                if span["font"].lower().find("bold") > -1 or span["font"].lower().find("heavy") > -1:
                                    if all(c.isupper() or not c.isalpha() for c in text):
                                        current_headings[0] = text
                                        current_headings[1] = None
                                        current_headings[2] = None
                                        current_headings[3] = None
                                        if text not in toc:
                                            toc[text] = {}
                                
                                # Level 2: Title Case, bold/heavy font
                                elif span["font"].lower().find("bold") > -1 or span["font"].lower().find("heavy") > -1:
                                    if text.istitle():
                                        current_headings[1] = text
                                        current_headings[2] = None
                                        current_headings[3] = None
                                        if current_headings[0]:
                                            if text not in toc[current_headings[0]]:
                                                toc[current_headings[0]][text] = {}

                                # Level 3: Title Case, regular font
                                elif text.istitle():
                                    current_headings[2] = text
                                    current_headings[3] = None
                                    if current_headings[0] and current_headings[1]:
                                        if text not in toc[current_headings[0]][current_headings[1]]:
                                            toc[current_headings[0]][current_headings[1]][text] = page_num - page_offset + 1
                                
                                # Level 4: Normal case
                                else:
                                    current_headings[3] = text
                                    if current_headings[0] and current_headings[1] and current_headings[2]:
                                        toc[current_headings[0]][current_headings[1]][current_headings[2]] = {text: page_num - page_offset + 1}
    return toc

class Exporter:
    def __init__(self, index: Dict, cross_references: Dict, toc: Dict, terms_only: bool = False, suppress_categories: List[str] = None, context_style: str = 'none'):
        self.cross_references = cross_references
        self.terms_only = terms_only
        self.suppress_categories = suppress_categories or []
        self.index = self._filter_index(index)
        self.toc = toc
        self.context_style = context_style

        if self.terms_only:
            self.case_law_references = {}
            self.statutory_references = {}
            self.subject_matter_index = {k: v for k, v in self.index.items()}
        else:
            self.case_law_references = {k: v for k, v in self.index.items() if 'case_law_references' in v}
            self.statutory_references = {k: v for k, v in self.index.items() if 'statutory_references' in v}
            self.subject_matter_index = {k: v for k, v in self.index.items() if 'case_law_references' not in v and 'statutory_references' not in v}

    def _filter_index(self, index: Dict) -> Dict:
        """Filter out suppressed categories from the index."""
        if not self.suppress_categories:
            return index
        
        filtered_index = defaultdict(lambda: defaultdict(list))
        for term, data in index.items():
            new_all_references = []
            # First, build the new filtered categories
            for category, pages in data.items():
                if category != 'all_references' and category not in self.suppress_categories:
                    filtered_index[term][category].extend(pages)
                    new_all_references.extend(pages)
            
            # If there are any remaining categories, add the new 'all_references'
            if filtered_index[term]:
                filtered_index[term]['all_references'] = new_all_references

        return {k: v for k, v in filtered_index.items() if v}

    def _format_entry(self, term, entries):
        if self.context_style == 'none' or len(entries) == 1:
            pages = sorted(list(set([entry[0] for entry in entries])))
            return f"{term}: {', '.join(map(str, pages))}"
        elif self.context_style == 'snippet':
            output = f"{term}:\n"
            for page, context, _ in sorted(entries):
                output += f"  - p. {page}: \"...{context}...\"\n"
            return output
        elif self.context_style == 'headings':
            output = f"{term}:\n"
            for page, _, headings in sorted(entries):
                output += f"  - p. {page}: {' > '.join(filter(None, headings))}\n"
            return output

    def to_text(self, include_subcategories: bool = True) -> str:
        """Generate formatted index with the new structure."""
        output = "COMPREHENSIVE LEGAL INDEX\n"
        output += "=" * 50 + "\n\n"

        # Table of Contents
        output += "TABLE OF CONTENTS\n"
        output += "-" * 50 + "\n"
        for l1, l2_dict in sorted(self.toc.items()):
            output += f"{l1}\n"
            for l2, l3_dict in sorted(l2_dict.items()):
                output += f"  {l2}\n"
                for l3, l4_val in sorted(l3_dict.items()):
                    if isinstance(l4_val, dict):
                        output += f"    {l3}: {list(l4_val.values())[0]}\n"
                        for l4, page in sorted(l4_val.items()):
                             output += f"      {l4}: {page}\n"
                    else:
                        output += f"    {l3}: {l4_val}\n"
        output += "\n"


        if not self.terms_only:
            # Case Law References
            if 'case_law_references' not in self.suppress_categories:
                output += "CASE LAW REFERENCES\n"
                output += "-" * 50 + "\n"
                for term, data in sorted(self.case_law_references.items()):
                    output += self._format_entry(term, data['all_references']) + "\n"
                output += "\n"

            # Statutory References
            if 'statutory_references' not in self.suppress_categories:
                output += "STATUTORY REFERENCES\n"
                output += "-" * 50 + "\n"
                for term, data in sorted(self.statutory_references.items()):
                    output += self._format_entry(term, data['all_references']) + "\n"
                output += "\n"

        # Index by Subject
        output += "INDEX BY SUBJECT\n"
        output += "-" * 50 + "\n"
        # Group terms by category
        subject_index = defaultdict(list)
        for term, data in sorted(self.subject_matter_index.items()):
            for category in data:
                if category != 'all_references' and category not in self.suppress_categories:
                    subject_index[category].append((term, data['all_references']))
        
        for category, terms in sorted(subject_index.items()):
            output += f"\n-- {category.replace('_', ' ').title()} --\n"
            for term, entries in sorted(terms):
                output += self._format_entry(term, entries) + "\n"
        output += "\n"

        # Alphabetical Index
        output += "ALPHABETICAL INDEX\n"
        output += "-" * 50 + "\n"
        
        index_source = self.subject_matter_index if self.terms_only else self.index
        for term, data in sorted(index_source.items()):
            output += self._format_entry(term, data['all_references']) + "\n"
        
        return output

    def to_json(self) -> str:
        """Convert index to JSON format with the new structure."""
        if self.terms_only:
            json_data = {
                'table_of_contents': self.toc,
                'subject_matter_index': {k: {cat: sorted(pages) for cat, pages in v.items()} for k, v in self.subject_matter_index.items()},
                '_cross_references': {k: sorted(list(v)) for k, v in self.cross_references.items()}
            }
        else:
            json_data = {
                'table_of_contents': self.toc,
                'case_law_references': {k: sorted(v['all_references']) for k, v in self.case_law_references.items()},
                'statutory_references': {k: sorted(v['all_references']) for k, v in self.statutory_references.items()},
                'subject_matter_index': {k: {cat: sorted(pages) for cat, pages in v.items()} for k, v in self.subject_matter_index.items()},
                '_cross_references': {k: sorted(list(v)) for k, v in self.cross_references.items()}
            }
        return json.dumps(json_data, indent=2, ensure_ascii=False)

    def to_pdf(self, path: str):
        """Generate a PDF version of the index."""
        doc = fitz.open()
        page = doc.new_page(width=595, height=842) # A4 size
        
        text = self.to_text()
        
        # Use a textbox for better layout control
        rect = fitz.Rect(50, 50, 545, 792)
        page.insert_textbox(rect, text, fontname="helv", fontsize=10)
        
        doc.save(path)
        doc.close()

    def to_docx(self, path: str, columns: int = 1):
        """Generate a .docx version of the index."""
        doc = Document()
        section = doc.sections[0]
        
        if columns > 1:
            sectPr = section._sectPr
            cols = sectPr.xpath('./w:cols')[0]
            cols.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}num', str(columns))

        doc.add_heading('Comprehensive Legal Index', 0)

        # Table of Contents
        doc.add_heading('Table of Contents', level=1)
        for l1, l2_dict in sorted(self.toc.items()):
            doc.add_paragraph(l1, style='Heading 2')
            for l2, l3_dict in sorted(l2_dict.items()):
                doc.add_paragraph(l2, style='Heading 3')
                for l3, l4_val in sorted(l3_dict.items()):
                    if isinstance(l4_val, dict):
                        doc.add_paragraph(f"{l3}: {list(l4_val.values())[0]}", style='Heading 4')
                        for l4, page in sorted(l4_val.items()):
                             doc.add_paragraph(f"{l4}: {page}", style='Normal')
                    else:
                        doc.add_paragraph(f"{l3}: {l4_val}", style='Heading 4')
        
        if not self.terms_only:
            # Case Law References
            if 'case_law_references' not in self.suppress_categories:
                doc.add_heading('Case Law References', level=1)
                for term, data in sorted(self.case_law_references.items()):
                    doc.add_paragraph(self._format_entry(term, data['all_references']))

            # Statutory References
            if 'statutory_references' not in self.suppress_categories:
                doc.add_heading('Statutory References', level=1)
                for term, data in sorted(self.statutory_references.items()):
                    doc.add_paragraph(self._format_entry(term, data['all_references']))

        # Index by Subject
        doc.add_heading('Index by Subject', level=1)
        subject_index = defaultdict(list)
        for term, data in sorted(self.subject_matter_index.items()):
            for category in data:
                if category != 'all_references' and category not in self.suppress_categories:
                    subject_index[category].append((term, data['all_references']))
        
        for category, terms in sorted(subject_index.items()):
            doc.add_heading(category.replace('_', ' ').title(), level=2)
            for term, entries in sorted(terms):
                doc.add_paragraph(self._format_entry(term, entries))

        # Alphabetical Index
        doc.add_heading('Alphabetical Index', level=1)
        index_source = self.subject_matter_index if self.terms_only else self.index
        for term, data in sorted(index_source.items()):
            doc.add_paragraph(self._format_entry(term, data['all_references']))
        
        doc.save(path)

    def to_csv(self, path: str):
        """Generate a CSV version of the index."""
        with open(path, 'w', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            writer.writerow(['Term', 'Category', 'Page', 'Context', 'Headings'])
            
            index_source = self.subject_matter_index if self.terms_only else self.index
            for term, data in sorted(index_source.items()):
                for cat, entries in data.items():
                    if entries and cat not in self.suppress_categories:
                        for page, context, headings in entries:
                            writer.writerow([term, cat, page, context, ' > '.join(filter(None, headings))])

    def to_xml(self, path: str):
        """Generate an XML version of the index."""
        root = etree.Element('LegalIndex')
        
        # Table of Contents
        toc_root = etree.SubElement(root, 'TableOfContents')
        for l1, l2_dict in sorted(self.toc.items()):
            l1_elem = etree.SubElement(toc_root, 'Heading1', name=l1)
            for l2, l3_dict in sorted(l2_dict.items()):
                l2_elem = etree.SubElement(l1_elem, 'Heading2', name=l2)
                for l3, l4_val in sorted(l3_dict.items()):
                    if isinstance(l4_val, dict):
                        l3_elem = etree.SubElement(l2_elem, 'Heading3', name=l3, page=str(list(l4_val.values())[0]))
                        for l4, page in sorted(l4_val.items()):
                            etree.SubElement(l3_elem, 'Heading4', name=l4, page=str(page))
                    else:
                        etree.SubElement(l2_elem, 'Heading3', name=l3, page=str(l4_val))

        if not self.terms_only:
            # Case Law
            if 'case_law_references' not in self.suppress_categories:
                case_law_root = etree.SubElement(root, 'CaseLawReferences')
                for term, data in sorted(self.case_law_references.items()):
                    ref_elem = etree.SubElement(case_law_root, 'Reference', name=term)
                    for page, context, _ in data['all_references']:
                        etree.SubElement(ref_elem, 'Occurrence', page=str(page)).text = context

            # Statutory
            if 'statutory_references' not in self.suppress_categories:
                statutory_root = etree.SubElement(root, 'StatutoryReferences')
                for term, data in sorted(self.statutory_references.items()):
                    ref_elem = etree.SubElement(statutory_root, 'Reference', name=term)
                    for page, context, _ in data['all_references']:
                        etree.SubElement(ref_elem, 'Occurrence', page=str(page)).text = context

        # Subject Matter
        subject_root = etree.SubElement(root, 'SubjectMatterIndex')
        for term, data in sorted(self.subject_matter_index.items()):
            term_element = etree.SubElement(subject_root, 'Term', name=term)
            for cat, entries in data.items():
                if cat not in self.suppress_categories:
                    cat_element = etree.SubElement(term_element, 'Category', name=cat)
                    for page, context, _ in entries:
                        etree.SubElement(cat_element, 'Occurrence', page=str(page)).text = context

        with open(path, 'wb') as f:
            f.write(etree.tostring(root, pretty_print=True, xml_declaration=True, encoding='UTF-8'))

    def to_markdown(self, path: str):
        """Generate a Markdown version of the index."""
        with open(path, 'w', encoding='utf-8') as f:
            f.write("# Comprehensive Legal Index\n\n")

            # Table of Contents
            f.write("## Table of Contents\n")
            for l1, l2_dict in sorted(self.toc.items()):
                f.write(f"### {l1}\n")
                for l2, l3_dict in sorted(l2_dict.items()):
                    f.write(f"#### {l2}\n")
                    for l3, l4_val in sorted(l3_dict.items()):
                        if isinstance(l4_val, dict):
                            f.write(f"- {l3}: {list(l4_val.values())[0]}\n")
                            for l4, page in sorted(l4_val.items()):
                                f.write(f"  - {l4}: {page}\n")
                        else:
                            f.write(f"- {l3}: {l4_val}\n")
            f.write("\n")


            if not self.terms_only:
                if 'case_law_references' not in self.suppress_categories:
                    f.write("## Case Law References\n")
                    for term, data in sorted(self.case_law_references.items()):
                        f.write(f"- {self._format_entry(term, data['all_references'])}\n")
                    f.write("\n")

                if 'statutory_references' not in self.suppress_categories:
                    f.write("## Statutory References\n")
                    for term, data in sorted(self.statutory_references.items()):
                        f.write(f"- {self._format_entry(term, data['all_references'])}\n")
                    f.write("\n")

            f.write("## Index by Subject\n")
            subject_index = defaultdict(list)
            for term, data in sorted(self.subject_matter_index.items()):
                for category in data:
                    if category != 'all_references' and category not in self.suppress_categories:
                        subject_index[category].append((term, data['all_references']))
            
            for category, terms in sorted(subject_index.items()):
                f.write(f"### {category.replace('_', ' ').title()}\n")
                for term, entries in sorted(terms):
                    f.write(f"- {self._format_entry(term, entries)}\n")
                f.write("\n")

            f.write("## Alphabetical Index\n")
            index_source = self.subject_matter_index if self.terms_only else self.index
            for term, data in sorted(index_source.items()):
                f.write(f"- {self._format_entry(term, data['all_references'])}\n")

def save_output(path: str, index: Dict, cross_references: Dict, toc: Dict, format: str, include_subcategories: bool = True, terms_only: bool = False, columns: int = 1, suppress_categories: List[str] = None, context_style: str = 'none'):
    """Save index output in the specified format."""
    exporter = Exporter(index, cross_references, toc, terms_only=terms_only, suppress_categories=suppress_categories, context_style=context_style)
    
    try:
        if format == 'text':
            with open(path, 'w', encoding='utf-8') as f:
                f.write(exporter.to_text(include_subcategories))
        elif format == 'json':
            with open(path, 'w', encoding='utf-8') as f:
                f.write(exporter.to_json())
        elif format == 'pdf':
            exporter.to_pdf(path)
        elif format == 'docx':
            exporter.to_docx(path, columns=columns)
        elif format == 'csv':
            exporter.to_csv(path)
        elif format == 'xml':
            exporter.to_xml(path)
        elif format == 'md':
            exporter.to_markdown(path)
        else:
            raise ValueError(f"Unsupported format: {format}")
            
        print(f"Index saved to: {path}")
        
    except Exception as e:
        print(f"Error saving output: {e}")

def get_statistics(index: Dict, page_content: Dict) -> Dict:
    """Get statistics about the indexed content."""
    stats = {
        'total_terms': len(index),
        'total_pages': len(page_content),
        'pages_with_content': sum(1 for text in page_content.values() if text.strip()),
    }
    
    category_counts = defaultdict(int)
    for term_data in index.values():
        for category in term_data.keys():
            if category != 'all_references':
                category_counts[category] += 1
    
    stats['terms_by_category'] = dict(category_counts)
    return stats
